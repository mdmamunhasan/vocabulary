from docx import Document


def split_list(a_list, part_by):
    """
    Split list into group of list
    :param a_list: input one dimensional list
    :param part_by: small list max length
    :return: 
    """
    s_list = []
    p_list = []
    for item in a_list:
        p_list.append(item.replace(':', ': '))
        if len(p_list) == part_by:
            s_list.append(p_list)
            p_list = []

    if len(p_list) > 0:
        s_list.append(p_list)

    return s_list


# Read content from text file
filename = 'vocabulary.txt'
file = open(filename, 'r')
content = file.read()
file.close()

# Format the vocabularies
vocabularies = []
for s in content.splitlines():
    vocabularies.append(s.strip())

vocabularies = list(reversed(vocabularies))
recordset = split_list(vocabularies, 2)

# Write the document
document = Document()

document.add_heading('Vocabulary', 0)

table = document.add_table(rows=len(recordset), cols=2)
for item in recordset:
    row_cells = table.add_row().cells

    col_count = len(item)
    row_cells[0].text = str(item[0])
    if col_count > 1:
        row_cells[1].text = str(item[1])
    print(item)

# document.add_page_break()

document.save('vocabulary.docx')
